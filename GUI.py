import os
from datetime import datetime
from docx import Document
from tkinter import Tk, Label, Entry, Button, filedialog, messagebox
from natsort import natsorted

class DocxGeneratorApp:
    def __init__(self, master):
        self.master = master
        master.title("Text to DOCX Converter")

        # Variables
        self.input_folder = ""
        self.output_folder = ""
        self.blacklist_file = ""
        self.blacklisted_files = []
        self.files = []

        # GUI Layout
        Label(master, text="Course Name and Code:").grid(row=0, column=0, sticky='w')
        self.course_entry = Entry(master, width=50)
        self.course_entry.grid(row=0, column=1, columnspan=2)

        Label(master, text="Submitted To:").grid(row=1, column=0, sticky='w')
        self.submitted_to_entry = Entry(master, width=50)
        self.submitted_to_entry.grid(row=1, column=1, columnspan=2)

        Label(master, text="Submitted By - Name:").grid(row=2, column=0, sticky='w')
        self.name_entry = Entry(master, width=50)
        self.name_entry.grid(row=2, column=1, columnspan=2)

        Label(master, text="ID:").grid(row=3, column=0, sticky='w')
        self.id_entry = Entry(master, width=50)
        self.id_entry.grid(row=3, column=1, columnspan=2)

        Label(master, text="Subject:").grid(row=4, column=0, sticky='w')
        self.subject_entry = Entry(master, width=50)
        self.subject_entry.grid(row=4, column=1, columnspan=2)

        Button(master, text="Select Input Folder", command=self.select_input_folder).grid(row=5, column=0, pady=5)
        Button(master, text="Select Output Folder", command=self.select_output_folder).grid(row=5, column=1)
        Button(master, text="Select Blacklist File", command=self.select_blacklist_file).grid(row=5, column=2)

        Button(master, text="Generate DOCX", command=self.generate_docx).grid(row=6, column=0, columnspan=3, pady=10)

    def select_input_folder(self):
        self.input_folder = filedialog.askdirectory(title="Select Input Folder")

    def select_output_folder(self):
        self.output_folder = filedialog.askdirectory(title="Select Output Folder")

    def select_blacklist_file(self):
        self.blacklist_file = filedialog.askopenfilename(title="Select Blacklist File")
        self.load_blacklist()

    def load_blacklist(self):
        if self.blacklist_file and os.path.exists(self.blacklist_file):
            with open(self.blacklist_file) as f:
                lines = [line.strip() for line in f if '#' not in line]
                self.blacklisted_files = [os.path.join(self.input_folder, line) for line in lines]

    def get_files(self, basepath):
        for filename in os.listdir(basepath):
            full_path = os.path.join(basepath, filename)
            if full_path in self.blacklisted_files:
                continue
            if os.path.isdir(full_path):
                self.get_files(full_path)
            else:
                self.files.append(full_path)

    def generate_docx(self):
        if not self.input_folder or not self.output_folder:
            messagebox.showerror("Error", "Please select both input and output folders.")
            return

        self.files = []
        self.get_files(self.input_folder)
        self.files = natsorted(self.files)

        doc = Document()
        doc.add_heading(self.course_entry.get(), 0)

        doc.add_paragraph('Submitted to:', style='Normal').runs[0].bold = True
        doc.add_paragraph(self.submitted_to_entry.get() + "\n")

        doc.add_paragraph('Submitted by:', style='Normal').runs[0].bold = True
        doc.add_paragraph(f"Name: {self.name_entry.get()}")
        doc.add_paragraph(f"ID: {self.id_entry.get()}")
        doc.add_paragraph(f"Subject: {self.subject_entry.get()}")
        doc.add_page_break()

        for fpath in self.files:
            with open(fpath, encoding="utf-8", errors='ignore') as f:
                doc.add_heading(os.path.relpath(fpath, self.input_folder), level=1)
                doc.add_paragraph(f.read())
                doc.add_page_break()

        now = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
        output_path = os.path.join(self.output_folder, f"output_{now}.docx")
        doc.save(output_path)

        messagebox.showinfo("Success", f"DOCX saved at:\n{output_path}")

# Run the GUI
if __name__ == "__main__":
    root = Tk()
    app = DocxGeneratorApp(root)
    root.mainloop()
