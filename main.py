from datetime import datetime
from docx import Document
import os
from natsort import natsorted

# ------------------------------
# Read blacklist
# ------------------------------
blacklisted_files = []

if os.path.exists("blacklist.txt"):
    with open("blacklist.txt") as f:
        blacklisted_files = [
            os.path.join("input", line.strip())
            for line in f
            if line.strip() and not line.strip().startswith("#")
        ]

files = []

# ------------------------------
# Collect files from input folder (and its subdirectories only)
# ------------------------------
def getfilesinfolder(basepath="input"):
    if not os.path.exists(basepath):
        print(f"Error: Directory '{basepath}' does not exist. Creating it now.")
        os.makedirs(basepath)
        print("Please put your files inside the 'input' directory and rerun.")
        return []

    # Walk through input/ and all its subdirectories
    for dirpath, dirnames, filenames in os.walk(basepath):
        # Process files in the current directory (including root)
        for filename in filenames:
            full_path = os.path.join(dirpath, filename)

            # Skip blacklisted files
            if full_path in blacklisted_files:
                continue

            # If the file is in the root folder, include it as well
            if dirpath == basepath:
                files.append(full_path)  # Add root files directly to list
            else:
                files.append(full_path)  # Add files from subdirectories as well

    return files

# ------------------------------
# Write Word document
# ------------------------------
def write():
    document = Document()

    # Title page
    document.add_heading('Title', 0)

    p = document.add_paragraph('Course name and code\n')

    p.add_run('Submitted to:\n').bold = True
    p.add_run('your information...\n\n\n')

    p.add_run('Submitted by:\n').bold = True
    p.add_run('Name:\n').bold = True
    p.add_run('ID:\n').bold = True
    p.add_run('Subject:\n').bold = True

    document.add_page_break()

    # Get and sort files from the input folder
    files_list = natsorted(getfilesinfolder())

    for filepath in files_list:
        # Generate heading like: lab1/a.py
        relative_path = os.path.relpath(filepath, "input")
        document.add_heading(relative_path, level=1)

        try:
            # Add file content with error handling for invalid characters
            with open(filepath, encoding="utf-8", errors="ignore") as f:
                document.add_paragraph(f.read())
        except Exception as e:
            print(f"Error reading file {filepath}: {e}")
            document.add_paragraph(f"Error reading file {filepath}")

        document.add_page_break()

    # Save file
    outputfile = "output.docx"
    document.save(outputfile)
    print(f"Saved as {outputfile}")


# ------------------------------
# Entry point
# ------------------------------
if __name__ == "__main__":
    write()

